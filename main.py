import re
from docx import shared
from pytesseract import image_to_string, pytesseract
from pathlib import Path
import docx
from PIL import Image
pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
topside_crop: int = 530
botside_crop: int = 240
content_file = docx.Document()
raw_screenshots = Path('./Screenshots').glob('*')
screenshot_ids = sorted([int(str(ss).split('_')[2]) for ss in raw_screenshots])
for i, img_id in enumerate(screenshot_ids):
	with Image.open(f'./Screenshots/Screenshot_20240520_{img_id}_Chrome.jpg') as img:
		width, height = img.size
		left: int = 100
		upper: int = topside_crop
		right: int = width - 100
		lower: int = height - botside_crop
		if upper < lower and left < right:
			cropped = img.crop((left, upper, right, lower))
			content = image_to_string(cropped, 'spa').replace('prisién', 'prisión').replace('ejecucién', 'ejecución').replace('detencién', 'detención')
			clean_content = re.sub(r'\We\s(?=[A-Z])', r'- ', content)
			if i > 0:
				content_file.add_paragraph('--------------------------------------')
			paragraph = content_file.add_paragraph()
			run = paragraph.add_run(clean_content)
			run.font.name = 'Arial'
			run.font.size = shared.Pt(14)
		else:
			print("Error: Las coordenadas de recorte no son válidas.")
content_file.save('ss_content.docx')